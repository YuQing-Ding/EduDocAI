"""

COPYRIGHT NOVA SCOTIA COMMUNITY COLLEGE - STRAIT AREA CAMPUS [ITGE]. ALL RIGHTS RESERVED.
PRODUCT MANAGER : DAVIS BOUDREAU
WRITTEN BY YUQING DING (SCOTT).

"""

import tkinter as tk
import os
import sys
from os import path
from tkinter import filedialog
from tkinter import PhotoImage
import openai
import json
import PyPDF2
import threading
from PIL import Image, ImageTk
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    base_path = getattr(sys, '_MEIPASS', path.dirname(path.abspath(__file__)))
    return path.join(base_path, relative_path)
selected_pdf_filename = ""
class PDFProcessor:
    def __init__(self):
        self.text = ""

    def extract_text(self, pdf_path):
        with open(pdf_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                self.text += page.extract_text()

    def get_text(self):
        return self.text

class ChatGPTInterface:
    def __init__(self, api_key):
        openai.api_key = api_key

    def generate_response(self, input_text):
        messages = [
            {"role": "system", "content": "An experienced lesson planner"},
            {"role": "user", "content": "Extract Course name and outcomes. Propose and analyze objectives professionally with pros and cons. Create optimal instructional program. \n" + input_text}
        ]

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=messages
        )

        return response.choices[0].message['content']

def get_api_key():
    openai_key = "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"  # Your OpenAI Key!
    return openai_key

def export_to_docx():
    if not response_text.get('1.0', tk.END).strip():
        tk.messagebox.showwarning("Warning!", "Text box is empty, cannot export to docx!")
        return

    doc = docx.Document()
    doc.add_heading('Here is the AI planned result', level=1)
    
    content = response_text.get('1.0', tk.END)
    content = content.strip()
    paragraphs = content.split('\n')

    for paragraph in paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.add_run(paragraph)

    doc.save(f"{selected_pdf_filename} - AI Planned.docx")
    docx_name = f"{selected_pdf_filename} - AI Planned.docx"
    tk.messagebox.showinfo("Success :)", f"The text has been successfully exported as a docx file.\n File name is: {docx_name}")

def show_loading_prompt():
    loading_window = tk.Toplevel(root)
    loading_window.title("ReadingEngineAI")
    loading_window.attributes('-topmost', True)  # Keep the loading window on top
    loading_window.overrideredirect(True)  # Remove the window decorations (no borders)
    loading_label = tk.Label(loading_window, text="ReadingEngine is reading and planning the document.\nPlease wait...")
    loading_label.pack(padx=10, pady=10)
    return loading_window

def process_pdf_and_generate_response():
    global selected_pdf_filename
    pdf_file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])

    if pdf_file_path:
        selected_pdf_filename = os.path.basename(pdf_file_path)
        selected_pdf_filename = selected_pdf_filename.split('_')[0]
        loading_window = show_loading_prompt()  # Show loading screen

        api_key = get_api_key()
        pdf_processor = PDFProcessor()
        pdf_processor.extract_text(pdf_file_path)
        extracted_text = pdf_processor.get_text()

        chat_gpt_interface = ChatGPTInterface(api_key)
        response = chat_gpt_interface.generate_response(extracted_text)

        response_text.config(state=tk.NORMAL)
        response_text.delete('1.0', tk.END)
        response_text.insert(tk.END, response)
        response_text.config(state=tk.DISABLED)

        loading_window.destroy()  # Kill loading screen!

def process_pdf_and_generate_response_with_loading():
    loading_thread = threading.Thread(target=process_pdf_and_generate_response)
    loading_thread.start()

# Create GUI window
root = tk.Tk()
root.title("ReaderEngineAI - Beta V0.01")
root.iconphoto(True, tk.PhotoImage(file=resource_path('icon.png')))
original_image = Image.open(resource_path('nscc_sa.png'))
resized_image = original_image.resize((128, 128))


# Create GUI components
choose_pdf_button = tk.Button(root, text="Select a PDF file", command=process_pdf_and_generate_response_with_loading)
export_button = tk.Button(root, text="Export the result as docx", command=export_to_docx)
response_text = tk.Text(root, wrap=tk.WORD, state=tk.DISABLED)
programmer_label = tk.Label(root, text="PROGRAMMER: YUQING (SCOTT) DING")
project_manager_label = tk.Label(root, text="PRODUCT MANAGER: DAVIS BOUDREAU")
copyright_label = tk.Label(root, text="© 2023 NSCC - NOVA SCOTIA COMMUNITY COLLEGE")
logo_image = ImageTk.PhotoImage(resized_image)
logo_label = tk.Label(root, image=logo_image)
# Layout GUI components
choose_pdf_button.pack(pady=10)
export_button.pack(pady=10)
response_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
programmer_label.pack(side="bottom", anchor="se", padx=10)
project_manager_label.pack(side="bottom", anchor="se", padx=10)
copyright_label.pack(side="bottom", anchor="se", padx=10, pady=10)
logo_label.pack(side="bottom", anchor="sw")

# GUI Main Loop
root.mainloop()